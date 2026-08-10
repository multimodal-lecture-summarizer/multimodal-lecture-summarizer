import io
import uuid
from typing import Optional
from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.core.exceptions import NotFoundException, ValidationException
from app.middleware.case_converter import CamelCaseAPIRoute
from app.schemas import BaseDTO, SummaryDTO
from app.api.deps import get_current_active_user
from app.models.user import User
from app.models.summary import Summary
from app.models.video import Video

router = APIRouter(route_class=CamelCaseAPIRoute)


@router.get(
    "/video/{video_id}",
    response_model=BaseDTO[SummaryDTO],
    summary="Get video summarization results",
    description="Retrieves the generated text summary, segmented chapters, and keyframes details for a video.",
)
def get_video_summary(
    video_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Retrieves AI summarization results for a processed video."""
    # Check ownership
    video = (
        db.query(Video)
        .filter(Video.video_id == video_id, Video.user_id == current_user.user_id)
        .first()
    )
    if not video:
        raise NotFoundException(message=f"Video with ID {video_id} not found")

    summary = db.query(Summary).filter(Summary.video_id == video_id).first()
    if not summary:
        raise NotFoundException(
            message="Summarization is still processing or has failed."
        )

    # In PostgreSQL, JSON fields might be automatically parsed into Python structures
    # We construct the response DTO manually to map the correct keyframe/chapter names.
    # Accept both camelCase (worker/sprint output) and snake_case for resilience.
    from app.schemas.summary import ChapterDTO, KeyframeDTO

    def _pick(obj: dict, *keys, default=None):
        for key in keys:
            if key in obj and obj[key] is not None:
                return obj[key]
        return default

    chapters = []
    for idx, c in enumerate(summary.chapters_json or []):
        if not isinstance(c, dict):
            continue
        start = float(_pick(c, "startTime", "start_time", "start_seconds", default=0.0) or 0.0)
        end = float(_pick(c, "endTime", "end_time", "end_seconds", default=start) or start)
        chapters.append(
            ChapterDTO(
                title=str(_pick(c, "title", default=f"Chapter {idx + 1}") or f"Chapter {idx + 1}"),
                start_time=start,
                end_time=end,
                summary=str(_pick(c, "summary", default="") or ""),
            )
        )

    keyframes = []
    for k in summary.keyframes_json or []:
        if not isinstance(k, dict):
            continue
        keyframes.append(
            KeyframeDTO(
                timestamp=float(_pick(k, "timestamp", default=0.0) or 0.0),
                image_url=str(_pick(k, "imageUrl", "image_url", default="") or ""),
                description=str(_pick(k, "description", "caption", default="") or ""),
                importance_score=float(
                    _pick(k, "importanceScore", "importance_score", default=0.5) or 0.5
                ),
            )
        )

    # Parse transcript_text to check if it's stored as a JSON list of segments
    import json
    import re
    parsed_segments = []
    clean_transcript_text = summary.transcript_text
    
    try:
        parsed = json.loads(summary.transcript_text)
        if isinstance(parsed, list):
            parsed_segments = parsed
            # Reconstruct clean text transcript for display/QA
            text_parts = []
            for seg in parsed:
                if seg.get("text") and seg.get("text") != "[Nhạc nền / Im lặng]":
                    text_parts.append(seg["text"])
            clean_transcript_text = " ".join(text_parts)
    except Exception:
        pass

    # If it is plain text, generate interpolated segments with word tokens for UI playback compatibility
    if not parsed_segments and clean_transcript_text:
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean_transcript_text) if s.strip()]
        total_dur = video.duration or 10.0
        sec_per_sentence = total_dur / len(sentences) if sentences else 10.0
        
        for idx, sentence in enumerate(sentences):
            s_start = idx * sec_per_sentence
            s_end = (idx + 1) * sec_per_sentence
            
            words_list = sentence.split()
            word_tokens = []
            if words_list:
                sec_per_word = sec_per_sentence / len(words_list)
                for w_idx, w in enumerate(words_list):
                    word_tokens.append({
                        "word": w,
                        "start": s_start + w_idx * sec_per_word,
                        "end": s_start + (w_idx + 1) * sec_per_word
                    })
                    
            parsed_segments.append({
                "speaker": "SPEAKER_01",
                "start": s_start,
                "end": s_end,
                "text": sentence,
                "words": word_tokens
            })

    dto = SummaryDTO(
        summary_id=summary.summary_id,
        video_id=summary.video_id,
        summary_text=summary.summary_text,
        chapters=chapters,
        keyframes=keyframes,
        transcript_text=clean_transcript_text,
        transcript_segments=parsed_segments,
        model_used=summary.model_used,
        processing_time=summary.processing_time,
    )

    return BaseDTO(
        success=True,
        data=dto,
        message="Summary results retrieved successfully",
    )


@router.get(
    "/video/{video_id}/export",
    summary="Export summary results to TXT, SRT, DOCX, or PDF files",
    description="Generates and streams a downloadable file of the summarization.",
)
def export_summary(
    video_id: uuid.UUID,
    format: str = Query("txt", description="File format to export: txt, srt, docx, pdf"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Generates downloadable summary files as StreamingResponse."""
    video = (
        db.query(Video)
        .filter(Video.video_id == video_id, Video.user_id == current_user.user_id)
        .first()
    )
    if not video:
        raise NotFoundException(message=f"Video with ID {video_id} not found")

    summary = db.query(Summary).filter(Summary.video_id == video_id).first()
    if not summary:
        raise NotFoundException(message="Summarization results not found.")

    format = format.lower().strip()
    if format not in ["txt", "srt", "docx", "pdf"]:
        raise ValidationException(
            message="Invalid format. Supported formats are: txt, srt, docx, pdf"
        )

    file_stream = io.BytesIO()
    
    # Parse transcript segments for export formats
    import json
    import re
    parsed_segments = []
    clean_transcript_text = summary.transcript_text
    
    try:
        parsed = json.loads(summary.transcript_text)
        if isinstance(parsed, list):
            parsed_segments = parsed
            text_parts = []
            for seg in parsed:
                if seg.get("text") and seg.get("text") != "[Nhạc nền / Im lặng]":
                    text_parts.append(seg["text"])
            clean_transcript_text = " ".join(text_parts)
    except Exception:
        pass

    if not parsed_segments and clean_transcript_text:
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean_transcript_text) if s.strip()]
        total_dur = video.duration or 10.0
        sec_per_sentence = total_dur / len(sentences) if sentences else 10.0
        
        for idx, sentence in enumerate(sentences):
            s_start = idx * sec_per_sentence
            s_end = (idx + 1) * sec_per_sentence
            parsed_segments.append({
                "start": s_start,
                "end": s_end,
                "text": sentence
            })

    def _chapter_export_fields(chapter: dict, idx: int) -> tuple[str, float, float, str]:
        title = chapter.get("title") or f"Chapter {idx}"
        start = chapter.get("startTime", chapter.get("start_time", chapter.get("start_seconds", 0)))
        end = chapter.get("endTime", chapter.get("end_time", chapter.get("end_seconds", start)))
        ch_summary = chapter.get("summary") or ""
        return title, start, end, ch_summary

    if format == "txt":
        content = (
            f"=== VIDEO SUMMARY REPORT ===\n"
            f"Video ID: {video_id}\n"
            f"Original URL: {video.original_url or 'N/A'}\n"
            f"Duration: {video.duration} seconds\n"
            f"Model Used: {summary.model_used}\n"
            f"Processing Time: {summary.processing_time}s\n"
            f"============================\n\n"
            f"SUMMARY:\n{summary.summary_text}\n\n"
            f"=== CHAPTERS ===\n"
        )
        for idx, c in enumerate(summary.chapters_json or [], 1):
            if not isinstance(c, dict):
                continue
            title, start, end, ch_summary = _chapter_export_fields(c, idx)
            content += f"Chapter {idx}: {title} ({start}s - {end}s)\n"
            content += f"Summary: {ch_summary}\n\n"

        content += f"=== TRANSCRIPT ===\n{clean_transcript_text or ''}\n"

        file_stream.write(content.encode("utf-8"))
        file_stream.seek(0)
        return StreamingResponse(
            file_stream,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=summary_{video_id}.txt"},
        )

    elif format == "srt":
        def format_srt_time(seconds: float) -> str:
            if not isinstance(seconds, (int, float)): return "00:00:00,000"
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            millis = int(round((seconds - int(seconds)) * 1000))
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

        srt_content = ""
        for i, seg in enumerate(parsed_segments, 1):
            start_str = format_srt_time(seg.get("start", 0))
            end_str = format_srt_time(seg.get("end", 0))
            text = seg.get("text", "")
            srt_content += f"{i}\n{start_str} --> {end_str}\n{text}\n\n"

        file_stream.write(srt_content.encode("utf-8"))
        file_stream.seek(0)
        return StreamingResponse(
            file_stream,
            media_type="text/srt",
            headers={"Content-Disposition": f"attachment; filename=subtitles_{video_id}.srt"},
        )

    elif format == "docx":
        try:
            from docx import Document
        except ImportError:
            raise ValidationException(message="DOCX export is not supported on this server.")
            
        doc = Document()
        doc.add_heading("Video Summary Report", 0)
        doc.add_paragraph(f"Video ID: {video_id}")
        doc.add_paragraph(f"Duration: {video.duration} seconds")
        
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary.summary_text)
        
        doc.add_heading("Chapters", level=1)
        for idx, c in enumerate(summary.chapters_json or [], 1):
            if not isinstance(c, dict):
                continue
            title, start, end, ch_summary = _chapter_export_fields(c, idx)
            doc.add_heading(f"Chapter {idx}: {title} ({start}s - {end}s)", level=2)
            doc.add_paragraph(ch_summary)
            
        doc.add_heading("Transcript", level=1)
        doc.add_paragraph(clean_transcript_text)
        
        doc.save(file_stream)
        file_stream.seek(0)
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=summary_{video_id}.docx"},
        )

    elif format == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            raise ValidationException(message="PDF export is not supported on this server.")
            
        import os
        
        pdf = FPDF()
        pdf.add_page()
        
        # Load font for Vietnamese support
        font_path = os.path.join(os.path.dirname(__file__), "..", "assets", "fonts", "Roboto-Regular.ttf")
        font_name = "Arial"
        if os.path.exists(font_path):
            pdf.add_font("Roboto", "", font_path, uni=True)
            font_name = "Roboto"
            
        pdf.set_font(font_name, size=16)
        pdf.cell(0, 10, "Video Summary Report", ln=True, align='C')
        pdf.set_font(font_name, size=12)
        pdf.cell(0, 10, f"Video ID: {video_id}", ln=True)
        pdf.cell(0, 10, f"Duration: {video.duration} seconds", ln=True)
        pdf.ln(5)
        
        pdf.set_font(font_name, size=14)
        pdf.cell(0, 10, "Summary", ln=True)
        pdf.set_font(font_name, size=11)
        pdf.multi_cell(0, 8, summary.summary_text)
        pdf.ln(5)
        
        pdf.set_font(font_name, size=14)
        pdf.cell(0, 10, "Chapters", ln=True)
        for idx, c in enumerate(summary.chapters_json or [], 1):
            if not isinstance(c, dict):
                continue
            title, start, end, ch_summary = _chapter_export_fields(c, idx)
            pdf.set_font(font_name, size=12)
            pdf.cell(0, 8, f"Chapter {idx}: {title} ({start}s - {end}s)", ln=True)
            pdf.set_font(font_name, size=11)
            pdf.multi_cell(0, 8, ch_summary)
            pdf.ln(3)
            
        pdf_bytes = pdf.output(dest='S')
        # fpdf2 dest='S' returns a bytearray
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin1')
            
        file_stream.write(pdf_bytes)
        file_stream.seek(0)
        return StreamingResponse(
            file_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=report_{video_id}.pdf"},
        )
